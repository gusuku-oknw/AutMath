import re
import os
import wx
import random
from fractions import Fraction
import sympy
from sympy import *
import datetime
import math
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.api import Document
from docx.opc.oxml import qn

from docx2pdf import convert

Name_Text = '＃__    名前_____________'
MathL = ''
MathL_answer = ''


class ThreeButtonEvent(wx.PyCommandEvent):
    def __init__(self, evtType, id):
        wx.PyCommandEvent.__init__(self, evtType, id)
        self.index = None

    def set_selected_index(self, index):
        self.index = index

    def get_selected_index(self):
        return self.index


myEVT_THREE_BUTTON = wx.NewEventType()
EVT_THREE_BUTTON = wx.PyEventBinder(myEVT_THREE_BUTTON, 1)


class ThreeButtonPanel(wx.Panel):
    def __init__(self, parent, id=-1):
        wx.Panel.__init__(self, parent, id)
        # Create widgets.
        self.button1 = wx.ToggleButton(self, label='生成')
        self.button2 = wx.ToggleButton(self, label='2')
        self.button3 = wx.ToggleButton(self, label='出力')
        self.button1.index = 0
        self.button2.index = 1
        self.button3.index = 2
        # Set event handlers.
        self.Bind(wx.EVT_TOGGLEBUTTON, self.on_toggle_button)
        # Set sizer.
        sizer = wx.BoxSizer(wx.HORIZONTAL)
        sizer.Add(self.button1)
        sizer.Add(self.button2)
        sizer.Add(self.button3)
        self.SetSizer(sizer)

    def on_toggle_button(self, evt):
        button = evt.GetEventObject()
        index = button.index
        if button.GetValue():
            if button != self.button1:
                self.button1.SetValue(False)
            if button != self.button2:
                self.button2.SetValue(False)
            if button != self.button3:
                self.button3.SetValue(False)
        else:
            index = -1
        # Raise event.
        evt = ThreeButtonEvent(myEVT_THREE_BUTTON, self.GetId())
        evt.set_selected_index(index)
        self.GetEventHandler().ProcessEvent(evt)


class AutMath:
    def __init__(self, percent, default_min=1, default_max=9):
        self.a = sympy.Symbol('a')
        self.b = sympy.Symbol('b')
        self.x = sympy.Symbol('x')
        self.y = sympy.Symbol('y')

        self.percent = percent
        print("default:" + str(percent))
        self.default_min = default_min
        self.default_max = default_max
        self.now_time()

    def AutMath_randint(self, rand_min=None, rand_max=None):
        rand = random.randint(self.default_min if rand_min is None else rand_min,
                              self.default_max if rand_max is None else rand_max)
        # print(rand)
        return rand

    def probability(self, percent=None):
        coin = random.random()
        # %で確率調整
        if coin * 100 < (self.percent if percent is None else percent):
            return True
        else:
            return False

    def carrying(self, value, mun, percent=None, numerical=None):
        # print(value)
        if mun < 1:
            return value
        # print(mun)
        if self.probability(percent=percent):
            value *= (random.randint(1, 10) if numerical is None else numerical)

        return self.carrying(value, mun - 1, percent=percent)

    def fraction_integer(self):
        x = random.randint(2 if self.default_min == 1 else self.default_min, self.default_max)
        while x == 0:
            x = random.randint(2 if self.default_min == 1 else self.default_min, self.default_max)
        y = random.randint(2 if self.default_min == 1 else self.default_min, self.default_max)
        while x % y == 0:
            y = random.randint(2 if self.default_min == 1 else self.default_min, self.default_max)

        print(Fraction(x, y))
        return Fraction(x, y)

    def search_rules(self, value, find_num=0):
        if '×' in value:
            r = value.rfind('×', find_num)
        elif '÷' in value:
            r = value.rfind('÷', find_num)
        elif '+' in value:
            r = value.rfind('+', find_num)
        elif '-' in value:
            r = value.rfind('-', find_num)
        else:
            r = -1
            print("不等号がありません")
        return r

    def replace_str2rules(self, value):
        val = str(value)
        val = val.replace('**', '^')
        val = val.replace('*', '')
        # for i in range(len(val)):
        #     print(val[i])
        #     if (val[i] == '×'):
        #         val.replace('×', '*')
        #     elif (val[i] == '÷'):
        #
        #     if not ('×' in value) or ('÷' in value):
        #         break
        # print("0")
        return val

    def replace_int2rules(self, rules):
        if rules == 1:
            rules = '+'
        elif rules == 2:
            rules = '-'
        elif rules == 3:
            rules = '×'
        elif rules == 4:
            rules = '÷'
        else:
            rules = rules

        return rules

    def problem_generation(self, question):
        if question is None:
            pass
        elif question == 1:
            value1 = self.AutMath_randint()
            value2 = self.AutMath_randint()

            value1 = self.carrying(value1, 1, percent=20)

            value1 = self.carrying(value1, 1, numerical=-1)
            value2 = self.carrying(value2, 1, numerical=-1)

            print(self.four_rules(value1, value2))

        elif question == 2:
            if self.probability():
                value1 = self.fraction_integer()
            else:
                value1 = Fraction(self.AutMath_randint(rand_min=2), 1)

            if value1 is not Fraction:
                value2 = self.fraction_integer()
            else:
                if self.probability():
                    value2 = self.fraction_integer()
                else:
                    value2 = Fraction(self.AutMath_randint(), 1)

            # value1 = self.carrying(value1, 1, percent=20)

            value1 = self.carrying(value1, 1, numerical=-1)
            value2 = self.carrying(value2, 1, numerical=-1)

            val, answer = self.four_rules(Fraction(value1), Fraction(value2))
            # Fraction(answer).limit_denominator(100)
            another_answer = Fraction(answer).limit_denominator(100)
            if answer != another_answer:
                answer = (answer, str(another_answer))
            print(val, answer)

        elif question == 3:
            value1 = self.AutMath_randint()
            value2 = self.AutMath_randint()
            value3 = self.AutMath_randint()

            value1 = self.carrying(value1, 1, percent=20)

            value1 = self.carrying(value1, 1, numerical=-1)
            value2 = self.carrying(value2, 1, numerical=-1)
            print("value1:{} value2:{} value3:{}".format(value1, value2, value3))
            val, answer = self.four_rules(value1, value2)
            print(self.four_rules(answer, value3, polynomial=val))

        elif question == 4:
            answer_sqrt = None

            value1 = self.AutMath_randint(rand_min=2, rand_max=7)
            value2 = self.AutMath_randint(rand_min=2, rand_max=7)
            while True:
                value_sqrt = self.AutMath_randint(rand_min=2, rand_max=5)
                value_sqrt_sub = self.AutMath_randint(rand_min=2, rand_max=5)
                if not (value_sqrt == 4) and (value_sqrt_sub == 4):
                    break

            value1 = self.carrying(value1, 1, numerical=-1)
            value2 = self.carrying(value2, 1, numerical=-1)

            # print("元の値{}、{}、{}".format(value1, value2, value_sqrt))

            val, answer = self.four_rules(value1, value2)
            # print("結果{}、{}".format(val, answer))
            if '÷' in val:
                r = val.rfind('÷')
                value1 = val[:r]

            if self.probability(percent=90):
                v = value_sqrt * (int(value1) ** 2)
                if v < 100:
                    val1_str = "√" + str(v)
                else:
                    val1_str = str(value1) + "√" + str(value_sqrt)
            else:
                val1_str = str(value1) + "√" + str(value_sqrt)
            if self.probability(percent=20):
                v = value_sqrt * (int(value2) ** 2)
                if v < 100:
                    val2_str = "√" + str(v)
                else:
                    val2_str = str(value2) + "√" + str(value_sqrt)
            else:
                val2_str = str(value2) + "√" + str(value_sqrt)
            r = self.search_rules(val)
            rules = val[r]

            if rules == '':
                answer_str = ''
            elif (rules == '+') or (rules == '-'):
                answer_str = str(answer) + "√" + str(value_sqrt)
            elif rules == '×':
                answer_str = str(answer * value_sqrt)
            elif rules == '÷':
                answer_str = str(answer)
            else:
                answer_str = ''
            val2_str = ("(" + str(val2_str) + ")" if val2_str[0] == '-' else str(val2_str))

            print("最終{}{}{}答え{}".format(val1_str, rules, val2_str, answer_str))

        elif question == 5:
            value1 = self.AutMath_randint()
            value2 = self.AutMath_randint()

            value1 = self.carrying(value1, 1, numerical=-1)
            value2 = self.carrying(value2, 1, numerical=-1)

            value1 *= self.a

            if self.probability(50):
                value1 *= self.a
            if self.probability(30):
                value1 *= self.b
            if self.probability(30):
                value1 *= self.b

            if self.probability(30):
                value2 *= self.a
            if self.probability(30):
                value2 *= self.a
            if self.probability(30):
                value2 *= self.b
            if self.probability(30):
                value2 *= self.b
            val, answer = self.four_rules(value1, value2)
            # print("val:{}".format(val))
            if self.probability(100):
                value3 = self.AutMath_randint()
                if self.probability(30):
                    value3 *= self.a
                if self.probability(30):
                    value3 *= self.a
                if self.probability(30):
                    value3 *= self.b
                if self.probability(30):
                    value3 *= self.b
                print(value1, value2, value3)
                val, answer = self.four_rules(answer, value3, polynomial=val)

            print("val:{}、answer:{}".format(val, answer))
            val = self.replace_str2rules(val)
            print(val)
            # pprint(sympify(str(val)))
        elif question == 6:
            # value1_0 ( value1_1 + value1_2 ) + value2_0 ( value2_1 + value2_2 )
            value1_0 = sympify(str(self.AutMath_randint(rand_max=3)))
            value1_1 = sympify(str(self.AutMath_randint(rand_max=3)))
            value1_2 = sympify(str(self.AutMath_randint(rand_max=3)))

            value1_0 = self.carrying(value1_0, 1, numerical=-1)
            value1_1 = self.carrying(value1_1, 1, numerical=-1)
            value1_2 = self.carrying(value1_2, 1, numerical=-1)

            value1_1 *= self.x
            value1_2 *= self.y

            value1 = value1_0 * value1_1 + value1_0 * value1_2
            value2_str = ''
            if self.probability(70):
                value2_0 = self.AutMath_randint(rand_max=3)
                value2_1 = self.AutMath_randint(rand_max=3)
                value2_2 = self.AutMath_randint(rand_max=3)

                value2_0 = self.carrying(value2_0, 1, numerical=-1)
                value2_1 = self.carrying(value2_1, 1, numerical=-1)
                value2_2 = self.carrying(value2_2, 1, numerical=-1)

                value2_1 *= self.x
                value2_2 *= self.y
                value2_str = str(value2_0 if value2_0 != 1 else '') + '(' + str(value2_1) + '+' + str(value2_2) + ')'
                value2 = value2_0 * value2_1 + value2_0 * value2_2
                val, answer = self.four_rules(sympify(value1), sympify(value2), rules='+')

                value1_str = str(value1_0 if value1_0 != 1 else '') + '(' + str(value1_1) + '+' + str(value1_2) + ')'
                val = value1_str + "+" + value2_str
            else:
                answer = expand(value1 ** 2)
                value1_0_str = str(value1_0 if value1_0 != 1 else '')
                value1_str = value1_0_str + '(' + str(value1_1) + '+' + str(value1_2) + ')' + '**2'
                val = value1_str

            print("val{}".format(val))
            print("answer:{}".format(answer))
            val = self.replace_str2rules(val)
            print(val)

    def four_rules(self, val1, val2, polynomial='', rules=''):
        poly_rules = ''
        rules = (random.randint(1, 4) if rules == '' else rules)

        rules = self.replace_int2rules(rules)

        poly_val1 = ''
        poly_val2 = ''
        # 与えられた値の文字列と結果を返す
        while polynomial != '':
            rules = random.randint(1, 4)

            rules = self.replace_int2rules(rules)

            if '-' in polynomial[0]:
                find_num = 1
            else:
                find_num = 0
            # print(find_num)
            poly_value = polynomial.replace("(", "").replace(")", "")

            r = self.search_rules(poly_value, find_num=find_num)

            poly_val1 = poly_value[:r]
            if poly_val1[-1] == '-':
                poly_val1 = poly_value[:r-1]
            poly_rules = poly_value[r]
            poly_val2 = poly_value[r + 1:]
            # print(poly_value[:r])
            # print(poly_val1, poly_val2)
            if ((rules == '+') or (rules == '-')) and (('×' in poly_rules) or ('÷' in poly_rules)):
                break

            elif ((rules == '×') or (rules == '÷')) and (('+' in poly_rules) or ('-' in poly_rules)):
                val1 = sympify(poly_val2)
                print(val1)
                break

        if rules == '':
            value = 0
            answer = 0

        elif rules == '+':
            val2_str = ("(" + str(val2) + ")" if str(val2)[0] == '-' else str(val2))
            value = ('+' + val2_str)
            answer = sympify(str(val1) +'+'+ str(val2))

        elif rules == '-':
            val2_str = ("(" + str(val2) + ")" if str(val2)[0] == '-' else str(val2))
            value = ('-' + val2_str)

            answer = sympify(str(val1) +'-'+ str(val2))

        elif rules == '×':
            val2_str = ("(" + str(val2) + ")" if str(val2)[0] == '-' else str(val2))
            value = ('×' + val2_str)
            if poly_rules != '':
                val, val1 = self.four_rules(poly_val1, val1, rules=poly_rules)
            answer = sympify(str(val1) +'*'+ str(val2))
            # int(val1) * val2

        elif rules == '÷':
            answer = val1

            # val1を求める
            val1 = sympify(str(val2) +'*'+ str(answer))

            if val1 is Fraction:
                answer = Fraction(val1)
                val2 = Fraction(val2)
                val1 = val2 * answer

            # print(val1)
            try:
                val1_abs = val1.subs([(self.a, 1), (self.b, 1), (self.x, 1), (self.y, 1)])
            except AttributeError or TypeError:
                val1_abs = float(val1)
            try:
                val2_abs = val2.subs([(self.a, 1), (self.b, 1), (self.x, 1), (self.y, 1)])
            except AttributeError or TypeError:
                val2_abs = float(val2)

            if val2_abs > val1_abs:
                val2, val1 = val1, val2

            val2_str = ("(" + str(val2) + ")" if str(val2)[0] == '-' else str(val2))
            value = ('÷' + val2_str)
            if poly_rules != '':
                val, val1 = self.four_rules(poly_val1, val1, rules=poly_rules)
                answer = sympify(str(val1) +'/'+ str(val2))

            poly_val2 = str(val1)

        else:
            value = None
            answer = None

        # print("value:{}answer:{}".format(value, answer))
        if polynomial != '':
            poly_val2 = ("(" + str(poly_val2) + ")" if str(poly_val2)[0] == '-' else str(poly_val2))
            return (poly_val1 + poly_rules + str(poly_val2) + value), answer
        return str(val1)+value, answer

    def now_time(self):
        dt_now = datetime.datetime.now()
        dt_name = str(dt_now.strftime('%Y.%m%d-%M%S'))
        return dt_name

    def mypath(self, file):
        cwd = os.path.abspath(file)
        # cwd += '/'
        print(cwd)
        return cwd

    # 問題docx生成のメソッド
    @staticmethod
    def add_docx(name):
        global MathL

        print(MathL)
        doc = Document()

        styles = doc.styles
        style = styles.add_style('Original-Style', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.size = Pt(15)
        font.name = 'BIZ UDゴシック'
        paragraph_format = style.paragraph_format

        p = doc.add_paragraph(MathL, style=style)

        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.save(name + ".docx")

        convert(name + ".docx", name + ".pdf")

    # 答えdocx生成のメソッド
    @staticmethod
    def add_docx_answer(name):
        global MathL_answer

        doc_answer = Document()

        styles = doc_answer.styles
        style = styles.add_style('Original-Style', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.size = Pt(15)
        font.name = 'BIZ UDゴシック'
        paragraph_format = style.paragraph_format

        p_answer = doc_answer.add_paragraph(MathL_answer, style=style)

        p_answer.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc_answer.save(name + ".docx")

        convert(name + ".docx", name + ".pdf")

        # convert()


class MyFrame(wx.Frame):
    def __init__(self):
        wx.Frame.__init__(self, None, -1, "Title", size=(300, 300))
        panel = ThreeButtonPanel(self)
        panel.Bind(EVT_THREE_BUTTON, self.OnThreeButton)

    @staticmethod
    def OnThreeButton(evt):
        global Name_Text
        global MathL
        global MathL_answer

        if evt.get_selected_index() == 0:
            max = 25
            MathL += Name_Text + '\n'
            MathL_answer += Name_Text + '\n'
            for i in range(max):
                math = AutMath()
                add_math = math.addition()

                MathL += '(' + str(1 + i).zfill(2) + ')' + add_math[0]
                MathL_answer += '(' + str(1 + i).zfill(2) + ')' + add_math[1]

                math = AutMath()
                sub_math = math.subtraction()
                MathL += '(' + str(26 + i) + ')' + sub_math[0]
                MathL_answer += '(' + str(26 + i) + ')' + sub_math[1]

                math = AutMath()
                mul_math = math.multiplication()
                MathL += '(' + str(51 + i) + ')' + mul_math[0]
                MathL_answer += '(' + str(51 + i) + ')' + mul_math[1]

                math = AutMath()
                div_math = math.division()
                MathL += '(' + str(76 + i) + ')' + div_math[0]
                MathL_answer += '(' + str(76 + i) + ')' + div_math[1]

                MathL += '\n'
                MathL_answer += '\n'

        if evt.get_selected_index() == 2:
            math = AutMath()

            # 問題のファイル名設定
            file_name = file_data.now_time()
            file_name += "_Q"
            file_name = file_data.mypath(file_name)

            # 答えのファイル名設定
            file_name_answer = file_data.now_time()
            file_name_answer += "_A"
            file_name_answer = file_data.mypath(file_name_answer)

            # 問題のdocxの生成
            math.add_docx(file_name)
            # 答えのdocxの生成
            math.add_docx_answer(file_name_answer)
            # 内容のリセット
            MathL = None
            MathL_answer = None

        print('Selected index =', evt.get_selected_index())


am = AutMath(50)
for i in range(50):
    am.problem_generation(6)

# print(vla)
#
# if __name__ == '__main__':
#     app = wx.PySimpleApp()
#     MyFrame().Show()
#     app.MainLoop()
