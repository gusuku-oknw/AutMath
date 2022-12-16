import os
import wx
import random
from fractions import Fraction
import datetime
import math
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.api import Document
from docx.opc.oxml import qn

from docx2pdf import convert

Name_Text = '＃__    名前_____________'
MathL = ''
MathL_answer = ''


class ThreeButtonEvent(wx.PyCommandEvent):
    def __init__(self, evtType, id):
        wx.PyCommandEvent.__init__(self, evtType, id)
        self.index = None

    def set_selected_index(self, index):
        self.index = index

    def get_selected_index(self):
        return self.index


myEVT_THREE_BUTTON = wx.NewEventType()
EVT_THREE_BUTTON = wx.PyEventBinder(myEVT_THREE_BUTTON, 1)


class ThreeButtonPanel(wx.Panel):
    def __init__(self, parent, id=-1):
        wx.Panel.__init__(self, parent, id)
        # Create widgets.
        self.button1 = wx.ToggleButton(self, label='生成')
        self.button2 = wx.ToggleButton(self, label='2')
        self.button3 = wx.ToggleButton(self, label='出力')
        self.button1.index = 0
        self.button2.index = 1
        self.button3.index = 2
        # Set event handlers.
        self.Bind(wx.EVT_TOGGLEBUTTON, self.on_toggle_button)
        # Set sizer.
        sizer = wx.BoxSizer(wx.HORIZONTAL)
        sizer.Add(self.button1)
        sizer.Add(self.button2)
        sizer.Add(self.button3)
        self.SetSizer(sizer)

    def on_toggle_button(self, evt):
        button = evt.GetEventObject()
        index = button.index
        if button.GetValue():
            if button != self.button1:
                self.button1.SetValue(False)
            if button != self.button2:
                self.button2.SetValue(False)
            if button != self.button3:
                self.button3.SetValue(False)
        else:
            index = -1
        # Raise event.
        evt = ThreeButtonEvent(myEVT_THREE_BUTTON, self.GetId())
        evt.set_selected_index(index)
        self.GetEventHandler().ProcessEvent(evt)


class AutMath:
    def __init__(self, percent):
        self.percent = percent
        print("default:"+str(percent))
        self.now_time()

    def probability(self, percent=None):
        coin = random.random()
        # %で確率調整
        if coin * 100 < (self.percent if percent is None else percent):
            return True
        else:
            return False

    def carrying(self, value, mun, percent=None, numerical=None):
        print(value)
        if mun < 1:
            return value
        # print(mun)
        if self.probability(percent=percent):
            value *= (random.randint(1, 10) if numerical is None else numerical)

        return self.carrying(value, mun - 1, percent=percent)

    def problem_generation(self):
        value1 = random.randint(1, 99)

        value2 = random.randint(1, 99)

        value1 = self.carrying(value1, 1)

        value1 = self.carrying(value1, 1, numerical=-1)

        value2 = self.carrying(value2, 1, numerical=-1)


    # 足し算のメソッド
    def addition(self, smallquestion):
        ad1 = random.randint(1, 99)
        ad2 = random.randint(1, 99)

        ad1 = self.carrying(ad1, 1)

        ad1 = self.carrying(ad1, 1, numerical=-1)

        ad2 = self.carrying(ad2, 1, numerical=-1)

        add = (str(ad1) + '+' + str("("+ad2+")" if ad2 < 0 else ad2) + '=  \t')
        ad3 = ad1 + ad2
        add_answer = ('=' + str(ad3) + '\t' + '\t')
        return add, add_answer

    # 引き算のメソッド
    def subtraction_integer(self):
        su1 = random.randint(1, 99)
        su2 = random.randint(1, 99)

        su1 = self.carrying(su1, 1)

        su1 = self.carrying(su1, 1, numerical=-1)

        su2 = self.carrying(su2, 1, numerical=-1)

        sub = (str(su1) + '-' + str("("+su2+")" if su2 < 0 else su2) + '=  \t')
        su3 = su1 - su2
        sub_answer = ('=' + str(su3) + '\t' + '\t')
        return sub, sub_answer

    # 掛け算のメソッド
    def multiplication_integer(self):
        mu1 = random.randint(1, 9)
        mu2 = random.randint(1, 9)

        mu1 = self.carrying(mu1, 1)

        mu2 = self.carrying(mu2, 1, percent=50)

        mu1 = self.carrying(mu1, 1, numerical=-1)

        mu2 = self.carrying(mu2, 1, numerical=-1)

        mul = (str(mu1) + '×' + str("("+mu2+")" if mu2 < 0 else mu2) + '=  \t')
        mul_answer = ('=' + str(mu1 * mu2) + '\t' + '\t')
        return mul, mul_answer

    # 割り算のメソッド
    def division_integer(self):
        di2 = random.randint(2, 9)
        answer = random.randint(2, 10)

        di2 = self.carrying(di2, 2)

        answer = self.carrying(answer, 2)

        di2 = self.carrying(di2, 1, numerical=-1)

        answer = self.carrying(answer, 1, numerical=-1)

        if di2 > answer:
            di2, answer = answer, di2

        # di1を求める
        di1 = di2 * answer

        sdiv = (str(di1) + '÷' + str("("+di2+")" if di2 < 0 else di2) + '=')
        sdiv_answer = ('=' + str(answer))
        return sdiv, sdiv_answer

    def now_time(self):
        dt_now = datetime.datetime.now()
        dt_name = str(dt_now.strftime('%Y.%m%d-%M%S'))
        return dt_name

    def mypath(self, file):
        cwd = os.path.abspath(file)
        # cwd += '/'
        print(cwd)
        return cwd

    # 問題docx生成のメソッド
    @staticmethod
    def add_docx(name):
        global MathL

        print(MathL)
        doc = Document()

        styles = doc.styles
        style = styles.add_style('Original-Style', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.size = Pt(15)
        font.name = 'BIZ UDゴシック'
        paragraph_format = style.paragraph_format

        p = doc.add_paragraph(MathL, style=style)

        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.save(name + ".docx")

        convert(name + ".docx", name + ".pdf")

    # 答えdocx生成のメソッド
    @staticmethod
    def add_docx_answer(name):
        global MathL_answer

        doc_answer = Document()

        styles = doc_answer.styles
        style = styles.add_style('Original-Style', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.size = Pt(15)
        font.name = 'BIZ UDゴシック'
        paragraph_format = style.paragraph_format

        p_answer = doc_answer.add_paragraph(MathL_answer, style=style)

        p_answer.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc_answer.save(name + ".docx")

        convert(name + ".docx", name + ".pdf")

        # convert()


class MyFrame(wx.Frame):
    def __init__(self):
        wx.Frame.__init__(self, None, -1, "Title", size=(300, 300))
        panel = ThreeButtonPanel(self)
        panel.Bind(EVT_THREE_BUTTON, self.OnThreeButton)

    @staticmethod
    def OnThreeButton(evt):
        global Name_Text
        global MathL
        global MathL_answer

        if evt.get_selected_index() == 0:
            max = 25
            MathL += Name_Text + '\n'
            MathL_answer += Name_Text + '\n'
            for i in range(max):
                math = AutMath()
                add_math = math.addition()

                MathL += '(' + str(1 + i).zfill(2) + ')' + add_math[0]
                MathL_answer += '(' + str(1 + i).zfill(2) + ')' + add_math[1]

                math = AutMath()
                sub_math = math.subtraction()
                MathL += '(' + str(26 + i) + ')' + sub_math[0]
                MathL_answer += '(' + str(26 + i) + ')' + sub_math[1]

                math = AutMath()
                mul_math = math.multiplication()
                MathL += '(' + str(51 + i) + ')' + mul_math[0]
                MathL_answer += '(' + str(51 + i) + ')' + mul_math[1]

                math = AutMath()
                div_math = math.division()
                MathL += '(' + str(76 + i) + ')' + div_math[0]
                MathL_answer += '(' + str(76 + i) + ')' + div_math[1]

                MathL += '\n'
                MathL_answer += '\n'

        if evt.get_selected_index() == 2:
            math = AutMath()

            # 問題のファイル名設定
            file_name = file_data.now_time()
            file_name += "_Q"
            file_name = file_data.mypath(file_name)

            # 答えのファイル名設定
            file_name_answer = file_data.now_time()
            file_name_answer += "_A"
            file_name_answer = file_data.mypath(file_name_answer)

            # 問題のdocxの生成
            math.add_docx(file_name)
            # 答えのdocxの生成
            math.add_docx_answer(file_name_answer)
            # 内容のリセット
            MathL = None
            MathL_answer = None

        print('Selected index =', evt.get_selected_index())


am = AutMath(50)
vla = am.carrying(4, 1, numerical=-1)
print(vla)
#
# if __name__ == '__main__':
#     app = wx.PySimpleApp()
#     MyFrame().Show()
#     app.MainLoop()
